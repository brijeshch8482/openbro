"""Word document tool — read, edit, and open .docx files.

Uses python-docx (install via openbro[office]).

Supports:
- read: extract all text
- info: paragraph / word / table counts
- find_replace: replace text across the document (returns replacement count)
- append: add a new paragraph at the end
- insert_after: find a marker paragraph and insert a new one after it
- open: launch the file in MS Word (or default associated app)
- list: list .docx files in a folder

Save policy: every edit is saved IN-PLACE by default. Pass `save_as=<path>`
to write a copy and leave the original untouched.
"""

from __future__ import annotations

import os
import platform
import subprocess
from pathlib import Path

from openbro.tools.base import BaseTool, RiskLevel

OFFICE_DEPS_HINT = (
    "Word/Excel deps not installed. Run: pip install 'openbro[office]' "
    "(or pip install python-docx openpyxl)"
)


def _resolve(path: str) -> Path:
    p = Path(path).expanduser()
    if not p.is_absolute():
        p = Path.cwd() / p
    return p.resolve()


def _ensure_docx(path: Path) -> str | None:
    if not path.exists():
        return f"File not found: {path}"
    if path.suffix.lower() != ".docx":
        return (
            f"Only .docx is supported (got {path.suffix}). "
            "Convert .doc → .docx in Word first (File → Save As)."
        )
    return None


class WordTool(BaseTool):
    name = "word"
    description = (
        "Read, edit, and open Microsoft Word (.docx) documents. "
        "Use 'open' to launch in Word, 'read' for text, "
        "'find_replace'/'append'/'insert_after' to edit, 'list' to find files."
    )
    risk = RiskLevel.MODERATE

    def schema(self) -> dict:
        return {
            "name": self.name,
            "description": self.description,
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": [
                            "open",
                            "read",
                            "info",
                            "find_replace",
                            "append",
                            "insert_after",
                            "list",
                        ],
                        "description": (
                            "Action: open=launch in Word, read=extract text, "
                            "info=stats, find_replace/append/insert_after=edit, "
                            "list=find .docx files in a folder."
                        ),
                    },
                    "file": {
                        "type": "string",
                        "description": "Absolute or relative path to the .docx file",
                    },
                    "find": {"type": "string", "description": "Text to find (find_replace)"},
                    "replace": {
                        "type": "string",
                        "description": "Replacement text (find_replace)",
                    },
                    "text": {
                        "type": "string",
                        "description": "Paragraph text to insert (append/insert_after)",
                    },
                    "after": {
                        "type": "string",
                        "description": (
                            "Marker paragraph to find for insert_after — first paragraph "
                            "containing this substring is the anchor"
                        ),
                    },
                    "save_as": {
                        "type": "string",
                        "description": ("Optional output path. Default: edit saves in-place."),
                    },
                    "folder": {
                        "type": "string",
                        "description": "Folder to search (for action=list). Default: current dir.",
                    },
                    "limit": {"type": "integer", "description": "Max chars to return for 'read'"},
                },
                "required": ["action"],
            },
        }

    def run(self, **kwargs) -> str:
        action = kwargs.get("action", "")

        if action == "list":
            return self._list(kwargs.get("folder", "."))

        if action == "open":
            return self._open(kwargs.get("file", ""))

        # All other actions need python-docx
        try:
            import docx  # noqa: F401
        except ImportError:
            return OFFICE_DEPS_HINT

        if action == "read":
            return self._read(kwargs.get("file", ""), kwargs.get("limit", 0))
        if action == "info":
            return self._info(kwargs.get("file", ""))
        if action == "find_replace":
            return self._find_replace(
                kwargs.get("file", ""),
                kwargs.get("find", ""),
                kwargs.get("replace", ""),
                kwargs.get("save_as"),
            )
        if action == "append":
            return self._append(
                kwargs.get("file", ""), kwargs.get("text", ""), kwargs.get("save_as")
            )
        if action == "insert_after":
            return self._insert_after(
                kwargs.get("file", ""),
                kwargs.get("after", ""),
                kwargs.get("text", ""),
                kwargs.get("save_as"),
            )
        return f"Unknown action: {action}"

    # ─── implementations ────────────────────────────────────

    def _list(self, folder: str) -> str:
        p = _resolve(folder)
        if not p.is_dir():
            return f"Folder not found: {p}"
        files = sorted(p.glob("**/*.docx"))
        if not files:
            return f"No .docx files in {p}"
        lines = [f"{f.relative_to(p)}  ({f.stat().st_size // 1024} KB)" for f in files[:50]]
        return "\n".join(lines) + (f"\n(+{len(files) - 50} more)" if len(files) > 50 else "")

    def _open(self, file_path: str) -> str:
        if not file_path:
            return "file is required."
        p = _resolve(file_path)
        if not p.exists():
            return f"File not found: {p}"
        try:
            if platform.system() == "Windows":
                os.startfile(str(p))  # type: ignore[attr-defined]
            elif platform.system() == "Darwin":
                subprocess.run(["open", str(p)], check=False)
            else:
                subprocess.run(["xdg-open", str(p)], check=False)
            return f"Opened in Word: {p}"
        except Exception as e:
            return f"Failed to open: {e}"

    def _read(self, file_path: str, limit: int = 0) -> str:
        from docx import Document

        p = _resolve(file_path)
        err = _ensure_docx(p)
        if err:
            return err
        try:
            doc = Document(str(p))
        except Exception as e:
            return f"Failed to open .docx: {e}"
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        if limit and limit > 0 and len(text) > limit:
            text = text[:limit] + f"\n... (+{len(text) - limit} more chars, raise 'limit' to see)"
        return text or "(document is empty)"

    def _info(self, file_path: str) -> str:
        from docx import Document

        p = _resolve(file_path)
        err = _ensure_docx(p)
        if err:
            return err
        try:
            doc = Document(str(p))
        except Exception as e:
            return f"Failed to open .docx: {e}"
        paras = [pa for pa in doc.paragraphs if pa.text.strip()]
        words = sum(len(pa.text.split()) for pa in paras)
        return (
            f"File: {p.name}\n"
            f"Size: {p.stat().st_size // 1024} KB\n"
            f"Paragraphs: {len(paras)}\n"
            f"Words: {words}\n"
            f"Tables: {len(doc.tables)}\n"
            f"Sections: {len(doc.sections)}"
        )

    def _find_replace(
        self,
        file_path: str,
        find: str,
        replace: str,
        save_as: str | None,
    ) -> str:
        from docx import Document

        if not find:
            return "'find' is required."
        p = _resolve(file_path)
        err = _ensure_docx(p)
        if err:
            return err
        try:
            doc = Document(str(p))
        except Exception as e:
            return f"Failed to open .docx: {e}"

        count = 0
        # Replace inside each paragraph's runs to preserve formatting where possible
        for para in doc.paragraphs:
            count += self._replace_in_paragraph(para, find, replace)
        # Also tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count += self._replace_in_paragraph(para, find, replace)

        if count == 0:
            return f"'{find}' not found — no changes made."

        out = _resolve(save_as) if save_as else p
        try:
            doc.save(str(out))
        except PermissionError:
            return f"Permission denied saving {out}. Is the file open in Word? Close it and retry."
        return f"Replaced {count} occurrence(s) of '{find}' → '{replace}'. Saved: {out}"

    def _append(
        self,
        file_path: str,
        text: str,
        save_as: str | None,
    ) -> str:
        from docx import Document

        if not text:
            return "'text' is required."
        p = _resolve(file_path)
        err = _ensure_docx(p)
        if err:
            return err
        try:
            doc = Document(str(p))
            doc.add_paragraph(text)
            out = _resolve(save_as) if save_as else p
            doc.save(str(out))
        except PermissionError:
            return f"Permission denied saving {out}. Close the file in Word first."
        except Exception as e:
            return f"Append failed: {e}"
        return f"Appended paragraph. Saved: {out}"

    def _insert_after(
        self,
        file_path: str,
        marker: str,
        text: str,
        save_as: str | None,
    ) -> str:
        from docx import Document

        if not marker or not text:
            return "Both 'after' (marker) and 'text' are required."
        p = _resolve(file_path)
        err = _ensure_docx(p)
        if err:
            return err
        try:
            doc = Document(str(p))
        except Exception as e:
            return f"Failed to open .docx: {e}"

        anchor = None
        for para in doc.paragraphs:
            if marker in para.text:
                anchor = para
                break
        if anchor is None:
            return f"Marker not found: '{marker}'"

        # Insert a new paragraph after the anchor by manipulating XML
        from docx.oxml.ns import qn  # type: ignore[import-not-found]

        new_p = anchor._p.makeelement(qn("w:p"), {})
        anchor._p.addnext(new_p)
        # Set its text via a fresh Paragraph wrapper
        from docx.text.paragraph import Paragraph  # type: ignore[import-not-found]

        new_para = Paragraph(new_p, anchor._parent)
        new_para.add_run(text)

        out = _resolve(save_as) if save_as else p
        try:
            doc.save(str(out))
        except PermissionError:
            return f"Permission denied saving {out}. Close the file in Word first."
        return f"Inserted after marker. Saved: {out}"

    # ─── helpers ────────────────────────────────────────────

    @staticmethod
    def _replace_in_paragraph(paragraph, find: str, replace: str) -> int:
        """Replace occurrences inside a paragraph, preserving runs where possible.

        Trade-off: when the find text spans multiple runs, we collapse them and
        lose run-level formatting. Common case (text inside one run) keeps style.
        """
        if find not in paragraph.text:
            return 0
        # Fast path: single run contains it
        for run in paragraph.runs:
            if find in run.text:
                count = run.text.count(find)
                run.text = run.text.replace(find, replace)
                return count
        # Multi-run path: collapse text, write back to first run, clear rest
        full = paragraph.text
        count = full.count(find)
        if count == 0:
            return 0
        new_text = full.replace(find, replace)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""
        return count
