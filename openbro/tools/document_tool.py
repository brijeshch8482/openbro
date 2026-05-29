"""Universal file reader — text out of any common file type.

Why this exists: the old setup only read `.docx` (word_tool), `.xlsx`
(excel_tool), and plain text (file_tool action='read'). When the user
asked the agent to read a PDF / image / audio / HTML the agent replied
"main sirf .docx files padh sakta hoon" — useless.

This tool dispatches by extension. Each backend dependency is
**optional**: if the user hasn't installed `openbro[docs]` we return a
clear "install X to read .Y" message instead of crashing.

Dispatch map (extension → backend):

  .pdf            -> pypdf text extraction; if extracted text is
                     effectively empty (scanned PDF), fall through to OCR
                     when `force_ocr=True` or when no text was found.
  .docx           -> python-docx
  .xlsx / .xls    -> openpyxl
  .csv / .tsv     -> stdlib csv
  .json           -> stdlib json (pretty-print)
  .yaml / .yml    -> PyYAML (already a core dep)
  .html / .htm    -> beautifulsoup4 (strip tags)
  .jpg .png .gif  -> pytesseract OCR (needs `tesseract` binary on PATH)
  .bmp .webp .tif
  .mp3 .wav .m4a  -> faster-whisper (already installed for voice STT)
  .ogg .flac .opus
  .zip            -> stdlib zipfile — list contents
  anything else   -> stdlib read_text(errors='replace')

All paths run through `resolve_user_path` so `~/Desktop/x.pdf` lands on
the OneDrive Desktop the user actually sees in Explorer.
"""

from __future__ import annotations

import csv as csv_mod
import io
import json as json_mod
import zipfile
from pathlib import Path

from openbro.tools.base import BaseTool, RiskLevel
from openbro.utils.paths import resolve_user_path

DOCS_DEPS_HINT = (
    "Optional doc-reading deps not installed. Install what you need:\n"
    "  PDF:    pip install pypdf\n"
    "  OCR:    pip install pytesseract Pillow  (+ tesseract binary on PATH)\n"
    "  HTML:   pip install beautifulsoup4\n"
    "  audio:  already covered by 'openbro[voice]' (faster-whisper)\n"
    "  all:    pip install 'openbro[docs]'"
)

PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xls"}
CSV_EXTS = {".csv", ".tsv"}
JSON_EXTS = {".json"}
YAML_EXTS = {".yaml", ".yml"}
HTML_EXTS = {".html", ".htm"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tif", ".tiff"}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".opus"}
ZIP_EXTS = {".zip"}
TEXT_EXTS = {
    ".txt",
    ".md",
    ".rst",
    ".log",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".go",
    ".rs",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".cs",
    ".rb",
    ".php",
    ".sh",
    ".ps1",
    ".bat",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".env",
    ".xml",
    ".sql",
}


class DocumentTool(BaseTool):
    name = "document"
    description = (
        "Read text out of ANY common file: PDF, image (OCR), audio (transcribe), "
        "HTML, CSV, JSON, YAML, Word, Excel, zip, code, text. Use this when "
        "the user wants to read/understand a file and you don't know the format "
        "in advance. Optional deps (pypdf, pytesseract, beautifulsoup4); the "
        "tool reports which to install if missing."
    )
    risk = RiskLevel.SAFE

    def schema(self) -> dict:
        return {
            "name": self.name,
            "description": self.description,
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "info"],
                        "description": (
                            "read = extract text from the file (any format), "
                            "info = report file type, size, and which backend "
                            "will be used without actually reading."
                        ),
                    },
                    "file": {
                        "type": "string",
                        "description": "Path to the file. Required.",
                    },
                    "limit": {
                        "type": "integer",
                        "description": (
                            "Max characters to return. Default 20000 — large enough "
                            "for normal docs, short enough to stay in context."
                        ),
                    },
                    "force_ocr": {
                        "type": "boolean",
                        "description": (
                            "For PDFs: skip text extraction and OCR every page. "
                            "Use when the PDF is scanned (text extraction "
                            "returns empty). Default false."
                        ),
                    },
                    "language": {
                        "type": "string",
                        "description": (
                            "For audio transcription, BCP-47 language code "
                            "(e.g. 'en', 'hi'). Default: auto-detect."
                        ),
                    },
                },
                "required": ["action", "file"],
            },
        }

    def run(self, **kwargs) -> str:
        action = kwargs.get("action", "read")
        file = kwargs.get("file", "")
        if not file:
            return "'file' is required."

        path = resolve_user_path(file)
        if not path.exists():
            return f"File not found: {path}"
        if path.is_dir():
            return f"That's a directory: {path}. Use file_ops action=list to list it."

        ext = path.suffix.lower()
        backend = self._backend_for(ext)

        if action == "info":
            size_kb = path.stat().st_size // 1024
            return (
                f"File: {path.name}\n"
                f"Path: {path}\n"
                f"Size: {size_kb} KB\n"
                f"Extension: {ext or '(none)'}\n"
                f"Backend: {backend}"
            )

        if action != "read":
            return f"Unknown action: {action}. Use 'read' or 'info'."

        limit = int(kwargs.get("limit") or 20000)
        try:
            if ext in PDF_EXTS:
                text = self._read_pdf(path, force_ocr=bool(kwargs.get("force_ocr", False)))
            elif ext in DOCX_EXTS:
                text = self._read_docx(path)
            elif ext in EXCEL_EXTS:
                text = self._read_excel(path)
            elif ext in CSV_EXTS:
                text = self._read_csv(path, ext)
            elif ext in JSON_EXTS:
                text = self._read_json(path)
            elif ext in YAML_EXTS:
                text = self._read_yaml(path)
            elif ext in HTML_EXTS:
                text = self._read_html(path)
            elif ext in IMAGE_EXTS:
                text = self._read_image(path)
            elif ext in AUDIO_EXTS:
                text = self._read_audio(path, language=kwargs.get("language"))
            elif ext in ZIP_EXTS:
                text = self._read_zip(path)
            else:
                # Text / code / unknown — best-effort text read.
                text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return f"Read failed ({backend}): {e}"

        if not text or not text.strip():
            return f"(no text extracted from {path.name})"

        if limit > 0 and len(text) > limit:
            text = text[:limit] + f"\n... (+{len(text) - limit} more chars, raise 'limit')"
        return text

    # ─── dispatch helper ───────────────────────────────────

    @staticmethod
    def _backend_for(ext: str) -> str:
        if ext in PDF_EXTS:
            return "pypdf (+ OCR fallback)"
        if ext in DOCX_EXTS:
            return "python-docx"
        if ext in EXCEL_EXTS:
            return "openpyxl"
        if ext in CSV_EXTS:
            return "csv (stdlib)"
        if ext in JSON_EXTS:
            return "json (stdlib)"
        if ext in YAML_EXTS:
            return "PyYAML"
        if ext in HTML_EXTS:
            return "beautifulsoup4"
        if ext in IMAGE_EXTS:
            return "pytesseract OCR"
        if ext in AUDIO_EXTS:
            return "faster-whisper"
        if ext in ZIP_EXTS:
            return "zipfile (stdlib)"
        if ext in TEXT_EXTS:
            return "text read"
        return "text read (fallback)"

    # ─── per-format readers ────────────────────────────────

    def _read_pdf(self, path: Path, force_ocr: bool = False) -> str:
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]
        except ImportError:
            return (
                "PDF reading needs pypdf. Install: pip install pypdf "
                "(or `pip install 'openbro[docs]'` for everything)."
            )

        text = ""
        if not force_ocr:
            try:
                reader = PdfReader(str(path))
                parts = []
                for page in reader.pages:
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        parts.append("")
                text = "\n\n".join(p for p in parts if p.strip())
            except Exception as e:
                return f"pypdf failed: {e}"

        # Scanned PDF: text extraction returns nothing. Try OCR if Tesseract
        # + pdf2image are available; otherwise tell the user how to enable it.
        if not text.strip() or force_ocr:
            ocr_text = self._ocr_pdf(path)
            if ocr_text is not None:
                return ocr_text
            if not text.strip():
                return (
                    "PDF text extraction returned nothing — looks like a scanned PDF. "
                    "Enable OCR: pip install pytesseract Pillow pdf2image  "
                    "(+ install the tesseract binary). Then retry with force_ocr=true."
                )
        return text

    def _ocr_pdf(self, path: Path) -> str | None:
        try:
            import pytesseract  # type: ignore[import-not-found]
            from pdf2image import convert_from_path  # type: ignore[import-not-found]
        except ImportError:
            return None
        try:
            images = convert_from_path(str(path))
        except Exception as e:
            return f"PDF→image conversion failed (needs poppler on PATH): {e}"
        parts = []
        for img in images:
            try:
                parts.append(pytesseract.image_to_string(img))
            except Exception as e:
                parts.append(f"[OCR error on page: {e}]")
        return "\n\n".join(parts)

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except ImportError:
            return "Word reading needs python-docx. Install: pip install 'openbro[office]'."
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text out of tables — common in resumes/reports.
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paras.append(row_text)
        return "\n".join(paras)

    def _read_excel(self, path: Path) -> str:
        try:
            from openpyxl import load_workbook  # type: ignore[import-not-found]
        except ImportError:
            return "Excel reading needs openpyxl. Install: pip install 'openbro[office]'."
        # data_only=True so formula cells return cached values, not '=A1+B1'.
        wb = load_workbook(str(path), data_only=True, read_only=True)
        out: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            out.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                out.append("\t".join(cells))
        return "\n".join(out)

    def _read_csv(self, path: Path, ext: str) -> str:
        delim = "\t" if ext == ".tsv" else ","
        with path.open(encoding="utf-8", errors="replace", newline="") as f:
            reader = csv_mod.reader(f, delimiter=delim)
            rows = list(reader)
        if not rows:
            return "(empty CSV)"
        # Pretty-print with header callout. Long files get summarized rather
        # than printing 100K lines — agent can ask python tool for analysis.
        out = [f"Rows: {len(rows)} | Cols: {len(rows[0])}"]
        out.append("\t".join(rows[0]))
        for r in rows[1:200]:
            out.append("\t".join(r))
        if len(rows) > 200:
            out.append(f"... (+{len(rows) - 200} more rows — use python tool for analysis)")
        return "\n".join(out)

    def _read_json(self, path: Path) -> str:
        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            data = json_mod.loads(raw)
            return json_mod.dumps(data, indent=2, ensure_ascii=False)
        except json_mod.JSONDecodeError:
            # Return raw so the agent can still see what's in the file.
            return raw

    def _read_yaml(self, path: Path) -> str:
        import yaml

        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            data = yaml.safe_load(raw)
            return yaml.dump(data, default_flow_style=False, allow_unicode=True)
        except yaml.YAMLError:
            return raw

    def _read_html(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup  # type: ignore[import-not-found]
        except ImportError:
            return (
                "HTML reading needs beautifulsoup4. Install: pip install beautifulsoup4 "
                "(or `pip install 'openbro[docs]'`)."
            )
        raw = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        # Collapse runs of blank lines to keep output readable.
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    def _read_image(self, path: Path) -> str:
        try:
            import pytesseract  # type: ignore[import-not-found]
            from PIL import Image  # type: ignore[import-not-found]
        except ImportError:
            return (
                "Image OCR needs pytesseract + Pillow. Install: pip install pytesseract Pillow "
                "(+ install the tesseract binary and put it on PATH). "
                "On Windows: https://github.com/UB-Mannheim/tesseract/wiki"
            )
        try:
            img = Image.open(str(path))
            text = pytesseract.image_to_string(img)
        except pytesseract.TesseractNotFoundError:
            return (
                "Tesseract binary not found on PATH. Install it: "
                "https://github.com/UB-Mannheim/tesseract/wiki (Windows) "
                "or `brew install tesseract` (mac) / `apt install tesseract-ocr` (Linux)."
            )
        return text or "(no text recognized in image)"

    def _read_audio(self, path: Path, language: str | None = None) -> str:
        try:
            from faster_whisper import WhisperModel  # type: ignore[import-not-found]
        except ImportError:
            return (
                "Audio transcription needs faster-whisper. Install: pip install 'openbro[voice]'."
            )
        # Reuse the same Whisper model size the voice listener uses. We
        # spin up a fresh instance here because the listener's STT may be
        # bound to its own thread; a separate instance is cheap on int8/cpu.
        model = WhisperModel("small", device="cpu", compute_type="int8")
        segments, _info = model.transcribe(str(path), language=language)
        text = " ".join(seg.text.strip() for seg in segments if seg.text)
        return text or "(no speech detected)"

    def _read_zip(self, path: Path) -> str:
        try:
            with zipfile.ZipFile(path) as zf:
                names = zf.namelist()
        except zipfile.BadZipFile:
            return f"Not a valid zip file: {path}"
        if not names:
            return "(empty zip)"
        buf = io.StringIO()
        buf.write(f"Archive: {path.name}  ({len(names)} entries)\n")
        for name in names[:200]:
            buf.write(f"  {name}\n")
        if len(names) > 200:
            buf.write(f"  ... (+{len(names) - 200} more)\n")
        buf.write("\n[hint: extract first, then use document.read on the file you want]")
        return buf.getvalue()
