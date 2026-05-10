"""Prepend a 'Current Implementation Status' section to the Vision doc.

Preserves the original vision content verbatim — just adds an up-to-date
status report at the top so the reader sees the as-built reality first,
then the full long-term vision.

Local-only doc (in .gitignore). Run with:
    python scripts/update_vision_doc.py
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

DOC = Path(__file__).resolve().parent.parent / "OpenBro_Vision_Document.docx"


# Each entry: (style, text). style values: 'h1', 'h2', 'h3', 'p', 'bullet', 'code'.
STATUS_BLOCKS: list[tuple[str, str]] = [
    ("h1", "Current Implementation Status — May 2026"),
    (
        "p",
        "Yeh section vision document ke top pe May 2026 me add kiya gaya — taaki reader "
        "current as-built state pehle dekhe, phir long-term vision. Original vision unchanged "
        "neeche se start hota hai.",
    ),
    (
        "p",
        "Status snapshot: OpenBro v1.0.0b1 — built, installed, running on developer machine. "
        "Major architectural pivot is complete: agent runs its own local LLM in-process via "
        "llama.cpp (no external Ollama daemon). Six LLM providers wired up with hybrid "
        "online/offline routing. Brain modules, MCP, custom desktop GUI, voice — all in.",
    ),
    # ─── Architecture summary ────────────────────────────────────────
    ("h2", "1. As-Built Architecture"),
    (
        "p",
        "OpenBro = full personal agent. The LLM is just the brain (text in, text out). "
        "Everything else — GUI, voice, memory, tools, automation, system control — is "
        "OpenBro's own code. No external runtimes (no Ollama daemon, no LM Studio, "
        "nothing to install separately).",
    ),
    ("h3", "Layer stack"),
    (
        "bullet",
        "Surfaces: customtkinter desktop GUI (default), terminal REPL (--cli), system tray "
        "with global hotkey Ctrl+Shift+B, Telegram bot, voice (mic + wake-word), MCP server.",
    ),
    (
        "bullet",
        "Agent core: openbro/core/agent.py — orchestrates conversations, tool calls, "
        "permission gates, language detection. Multi-role pipeline (planner → executor → "
        "verifier) for complex requests via openbro/core/multi_role.py.",
    ),
    (
        "bullet",
        "Brain: openbro/brain/* — persistent memory (semantic + keyword fallback), skill "
        "registry, reflection (learns from each interaction), world snapshot (PC facts), "
        "self-coder (LLM writes Python on demand and saves as a skill), community-manifest "
        "updater, daily LLM-upgrade probe.",
    ),
    (
        "bullet",
        "Local LLM: openbro/llm/local_engine.py wraps llama-cpp-python (Apache 2.0 "
        "llama.cpp). Models are GGUF files on disk; downloaded directly from HuggingFace "
        "via httpx streaming (no XET negotiation stall). In-process, no daemon.",
    ),
    (
        "bullet",
        "Cloud LLMs: Anthropic / OpenAI / Groq / Google / DeepSeek — each with its own "
        "provider class implementing the openbro.llm.base.LLMProvider interface. "
        "openbro/llm/router.py picks one from config; openbro/llm/auto_select.py probes "
        "what's available and ranks by capability.",
    ),
    (
        "bullet",
        "Tools: openbro/tools/* — file ops, command runner, web fetch, app launcher, "
        "memory tool, document editing (Word/Excel via python-docx/openpyxl). Each tool "
        "tagged safe / moderate / dangerous, gated by openbro/core/permissions.py.",
    ),
    (
        "bullet",
        "MCP: openbro/mcp/{client,server}.py — JSON-RPC over stdio. OpenBro can both "
        "consume external MCP servers (filesystem, github, sqlite, time, fetch) and "
        "expose itself as an MCP server (so Claude Desktop etc. can call OpenBro tools).",
    ),
    (
        "bullet",
        "Storage: respects user's chosen drive (wizard step 1). Memory DB, skills, model "
        "files, audit log, brain manifest — all in storage.base_dir / models_dir.",
    ),
    # ─── LLM strategy ──────────────────────────────────────────────────
    ("h2", "2. LLM Strategy — Hybrid Online + Offline"),
    (
        "p",
        "Six providers visible in the wizard's LLM step. User picks one at setup; can "
        "switch any time via 'model switch <alias>'. Daily auto-check probes vendor APIs "
        "for newer model releases.",
    ),
    ("h3", "Wizard step 2 — provider catalogue"),
    ("code", "1. Groq (FREE)            — llama-3.3-70b-versatile (recommended)"),
    ("code", "2. Google Gemini (FREE)   — gemini-1.5-flash"),
    ("code", "3. OpenAI (FREE-TRIAL)    — gpt-4o-mini"),
    ("code", "4. Anthropic Claude (PAID) — claude-sonnet-4-20250514"),
    ("code", "5. DeepSeek (CHEAP)       — deepseek-chat"),
    ("code", "6. Local (offline) (OFFLINE) — llama3.1:8b via llama.cpp"),
    ("h3", "Why llama.cpp instead of Ollama"),
    (
        "p",
        "Ollama was the placeholder during early development but it's a separate daemon, "
        "needs its own install, runs in the background, and adds an HTTP hop. We use the "
        "same engine (llama.cpp is what Ollama wraps internally) directly via Python "
        "bindings — no daemon, no HTTP hop, ~10-20% lower latency, simpler deps.",
    ),
    (
        "p",
        "User reaction during the migration captured this clearly: 'hamare agent ka khud "
        "ka GUI hona chahiye, llm ko handle krna ye sab khud krega — hame bs llm chahiye "
        "(brain). LLM only brain, uske aage GUI, handling, thinking, memory all things "
        "hamara agent karega.' That's the architecture now.",
    ),
    # ─── Local model catalogue ────────────────────────────────────────
    ("h2", "3. Local Model Catalogue (GGUF, non-Chinese vendors only)"),
    (
        "p",
        "Ten curated models from Meta, Mistral AI (France), Microsoft, and Google. All "
        "Q4_K_M quantization (best size/quality trade-off for most users). All hosted on "
        "HuggingFace under bartowski/* namespace (community-standard quantizations).",
    ),
    ("code", "Model            Size     RAM     Vendor / Strength"),
    ("code", "─────────────────────────────────────────────────────────────────"),
    ("code", "llama3.1:8b      4.9 GB   8 GB    Meta — best agent (recommended)"),
    ("code", "llama3.2:3b      2.0 GB   4 GB    Meta — smaller agent"),
    ("code", "llama3.2:1b      0.8 GB   2 GB    Meta — low-end PCs"),
    ("code", "mistral:7b       4.4 GB   8 GB    Mistral — reliable all-rounder"),
    ("code", "mistral-nemo     7.5 GB  12 GB    Mistral — 128K context"),
    ("code", "codestral:22b   13 GB   16 GB    Mistral — code specialist"),
    ("code", "phi3:mini        2.4 GB   4 GB    Microsoft — tiny but smart"),
    ("code", "phi3:medium      8.6 GB  16 GB    Microsoft — strong reasoning"),
    ("code", "gemma2:2b        1.7 GB   4 GB    Google — lightweight chat"),
    ("code", "gemma2:9b        5.8 GB   8 GB    Google — strong general chat"),
    (
        "p",
        "Project rule excludes Chinese-vendor models (Qwen, Yi, DeepSeek's open-weight) "
        "from the local catalogue, per user's requirement during the migration.",
    ),
    # ─── CLI commands ─────────────────────────────────────────────────
    ("h2", "4. CLI Commands (current surface)"),
    ("code", "openbro                       # default: launches desktop GUI"),
    ("code", "openbro --cli                 # terminal REPL instead"),
    ("code", "openbro --setup               # re-run first-run wizard"),
    ("code", "openbro --voice               # voice-only mode (mic + TTS)"),
    ("code", "openbro --tray                # system tray + global hotkey"),
    ("code", "openbro --telegram            # run as Telegram bot"),
    ("code", "openbro --mcp-server          # expose self as MCP server (stdio)"),
    ("code", "openbro --offline             # force provider = local"),
    ("code", "openbro -p <provider> -m <model>  # override LLM for this run"),
    ("code", ""),
    ("code", "openbro model download llama3.1:8b   # fetch a GGUF from HuggingFace"),
    ("code", "openbro model import path/file.gguf  # USB-transferred file (air-gapped)"),
    ("code", "openbro model list                   # installed + catalogue"),
    ("code", "openbro model remove <name>          # free disk"),
    # ─── Brain ─────────────────────────────────────────────────────────
    ("h2", "5. Brain Modules (persistent learning)"),
    (
        "bullet",
        "Brain.load() — singleton loaded at agent startup. Persists at storage.base_dir/"
        "brain/ as a manifest + sqlite + JSON files. Portable: 'brain export' / 'brain "
        "import' tar.gz for backup / migration.",
    ),
    (
        "bullet",
        "SemanticMemory — sentence-transformers embeddings + sqlite. Falls back to "
        "keyword search if sentence-transformers isn't installed. Compacts to recent N "
        "days on demand.",
    ),
    (
        "bullet",
        "SkillRegistry — learned skills with usage counts, success/failure rates, and "
        "trigger phrases. Skills live as Python files; the agent can write new ones via "
        "self_coder.py and add them at runtime.",
    ),
    (
        "bullet",
        "Reflector — after each interaction, classifies the signal (positive / "
        "negative / neutral), updates skill scores, and records a learning event. Drives "
        "long-term improvement.",
    ),
    (
        "bullet",
        "World — captures PC facts at startup (OS, hostname, paths, installed apps, "
        "online/offline). Refreshes every 6h. Injected into the system prompt so the "
        "model knows the user's environment without re-asking.",
    ),
    (
        "bullet",
        "Updater — daily check (24h cooldown) probes vendor APIs (Anthropic, OpenAI, "
        "Groq, Google) for newer models the user has access to. If a meaningfully better "
        "one exists, suggests it once.",
    ),
    # ─── Install flow ─────────────────────────────────────────────────
    ("h2", "6. One-Liner Install (Windows / macOS / Linux)"),
    ("code", "# Windows PowerShell (SHA-pinned for cache-bust safety):"),
    (
        "code",
        "$sha=(iwr -useb 'https://api.github.com/repos/brijeshch8482/openbro/commits/main"
        "'|ConvertFrom-Json).sha; iwr -useb \"https://raw.githubusercontent.com/brijeshch"
        "8482/openbro/$sha/scripts/install.ps1\" | iex",
    ),
    ("code", ""),
    ("code", "# Mac / Linux:"),
    ("code", "curl -fsSL https://raw.githubusercontent.com/brijeshch8482/openbro/main/scripts/install.sh | bash"),
    (
        "p",
        "What it does: detects Python (auto-installs 3.12 if current is too new for voice "
        "wheels), checks Node.js (for MCP servers via npx), pip-installs openbro[all,voice] "
        "from GitHub HEAD with --extra-index-url for llama-cpp-python's wheel index, then "
        "launches the wizard.",
    ),
    # ─── Bug history ──────────────────────────────────────────────────
    ("h2", "7. Notable Bugs Resolved During the Migration (May 2026)"),
    (
        "bullet",
        "PowerShell EAP=Stop crashing on pip stderr — fixed via local EAP=Continue + "
        "PSNativeCommandUseErrorActionPreference toggle in Invoke-Pip.",
    ),
    (
        "bullet",
        "voice deps source-compile killing the shell on Python 3.14 — installer now "
        "auto-installs Python 3.12 alongside and uses that for OpenBro; --only-binary=:all: "
        "guards against future source-build crashes.",
    ),
    (
        "bullet",
        "GitHub CDN serving stale install.ps1 due to ISP transparent caches stripping "
        "query strings — switched to SHA-pinned URLs via the GitHub API.",
    ),
    (
        "bullet",
        "pip's '#egg=name[extra]' fragment rejected on pip >= 23 — switched to PEP 508 "
        "direct-URL form: 'name[extra] @ git+https://...'.",
    ),
    (
        "bullet",
        "llama-cpp-python has NO wheels on PyPI (source-only); pip's source-compile hit "
        "Windows long-path limit on a deeply-nested vendor file — added "
        "--extra-index-url https://abetlen.github.io/llama-cpp-python/whl/cpu so wheels "
        "are pulled from the project's own GitHub Pages index.",
    ),
    (
        "bullet",
        "Invoke-Pip was returning the entire pip output array along with the int exit "
        "code; caller's '$exit -ne 0' was always true → silent PyPI fallback to stale "
        "1.0.0-beta. Fixed via 'Out-Host' so function returns only the int.",
    ),
    (
        "bullet",
        "openbro.__version__ was hardcoded as '1.0.0-beta' but pyproject.toml said "
        "'1.0.0b1'; verify step kept reporting stale version after a successful upgrade. "
        "Switched to importlib.metadata so __version__ always matches what pip installed.",
    ),
    (
        "bullet",
        "pip 'install --upgrade git+url' silently skipped reinstalling openbro itself "
        "when version comparison decided the existing copy was acceptable. Force-reinstall "
        "(--no-deps to keep cached deps) added when the installer detects an existing "
        "install.",
    ),
    (
        "bullet",
        "huggingface_hub.hf_hub_download stalled at '0%' for 30-60 sec on first download "
        "while XET protocol negotiated chunk manifests — switched to direct httpx.stream "
        "from the public 'resolve/main' URL, with .part file + atomic rename + Range-header "
        "resume.",
    ),
    # ─── Status of original 10-week roadmap ─────────────────────────
    ("h2", "8. Status vs the Original 10-Week Roadmap"),
    (
        "p",
        "The original vision below proposed a phased v0.1 → v1.0 ramp over 10 weeks. "
        "Reality: built all advanced features at once on main per user direction "
        "('phased rollout nahi, sab ek saath bana — ham advance test krenge'). Current "
        "tag 1.0.0b1 maps roughly to:",
    ),
    (
        "bullet",
        "v0.1 Foundation — DONE: agent core, six providers, REPL, GUI, system tools.",
    ),
    (
        "bullet",
        "v0.2 Tool Expansion — DONE: file/command/web/app/document tools, permission "
        "gating, audit log.",
    ),
    (
        "bullet",
        "v0.3 Telegram + Memory — DONE: Telegram bot wired, three-tier memory (working / "
        "long-term sqlite / semantic via sentence-transformers).",
    ),
    (
        "bullet",
        "v0.4 Skills System — DONE plus self-coder (LLM writes new Python skills on "
        "demand) and BroHub-style community manifest pull.",
    ),
    (
        "bullet",
        "v0.5 Voice Layer — DONE: faster-whisper STT (offline), Edge-TTS / pyttsx3, wake "
        "word detection. Cloud STT (Groq Whisper) available as opt-in fallback.",
    ),
    (
        "bullet",
        "MCP integration — DONE (was post-v1.0 in original plan): full client + server, "
        "JSON-RPC over stdio, five servers in wizard catalogue.",
    ),
    (
        "bullet",
        "System tray + global hotkey — DONE (was v1.5 in original plan): pystray + "
        "pynput, Ctrl+Shift+B opens main window.",
    ),
    (
        "bullet",
        "Custom desktop GUI — DONE (was post-v1.0 in original plan): customtkinter, "
        "voice button, slash commands, CLI streaming.",
    ),
    (
        "bullet",
        "PENDING: real-world long-form testing of the full advance v1, mobile app (v2), "
        "BroHub skill marketplace UI, multi-machine brain sync.",
    ),
    # ─── End of new section ─────────────────────────────────────────
    ("h2", "9. Where the Original Vision Document Begins"),
    (
        "p",
        "Everything below this line is the original April 2026 vision document, preserved "
        "verbatim. Read it for the long-term motivation, principles, target audience, and "
        "philosophy. The status above tells you what's actually built today; the vision "
        "tells you why.",
    ),
]


def style_run(run, *, bold=False, mono=False, color=None, size=None):
    if bold:
        run.bold = True
    if mono:
        run.font.name = "Consolas"
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)


def build_status_xml(doc: Document):
    """Create the status paragraphs in a fresh temporary doc, then return their XML
    so we can splice them into the start of the real document body."""
    tmp = Document()
    for kind, text in STATUS_BLOCKS:
        if kind == "h1":
            tmp.add_heading(text, level=1)
        elif kind == "h2":
            tmp.add_heading(text, level=2)
        elif kind == "h3":
            tmp.add_heading(text, level=3)
        elif kind == "bullet":
            p = tmp.add_paragraph(text, style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
        elif kind == "code":
            p = tmp.add_paragraph()
            run = p.add_run(text)
            style_run(run, mono=True, size=10, color=(0x33, 0x33, 0x33))
            p.paragraph_format.space_after = Pt(0)
        else:  # 'p'
            tmp.add_paragraph(text)

    # Page break after the status section so the original vision starts on a fresh page
    p = tmp.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    return [deepcopy(p._p) for p in tmp.paragraphs]


def main():
    if not DOC.exists():
        print(f"Doc not found: {DOC}")
        return

    doc = Document(str(DOC))

    # If we've already prepended (idempotent run), skip.
    first_heading = next(
        (p for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")),
        None,
    )
    if first_heading and "Current Implementation Status" in first_heading.text:
        print("Status section already present — nothing to do.")
        return

    body = doc.element.body
    new_paras = build_status_xml(doc)

    # Insert before the original first paragraph. addprevious() places the new
    # element directly before the target — and since insert_before is a fixed
    # reference to the ORIGINAL first paragraph, iterating new_paras in forward
    # order places them in the correct order:
    #   [new_para_1, ORIG_FIRST] → [new_para_1, new_para_2, ORIG_FIRST] → ...
    insert_before = body[0]
    for para_xml in new_paras:
        insert_before.addprevious(para_xml)

    out = DOC
    doc.save(str(out))

    # Quick sanity check
    saved = Document(str(out))
    para_count = len([p for p in saved.paragraphs if p.text.strip()])
    print(f"Saved: {out}")
    print(f"Total non-empty paragraphs: {para_count}")
    print(f"First heading now: {saved.paragraphs[0].text[:80]}")


if __name__ == "__main__":
    main()
